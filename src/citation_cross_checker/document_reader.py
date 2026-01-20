"""Document reader utilities for various file formats."""

from pathlib import Path
from typing import Union


class DocumentReader:
    """Reads text from various document formats."""

    @staticmethod
    def read_file(file_path: Union[str, Path]) -> str:
        """
        Read text from a file.

        Supports:
        - Plain text files (.txt)
        - Markdown files (.md)
        - Word documents (.docx)
        - PDF files (.pdf)

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is unsupported
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix in ['.txt', '.md', '.markdown', '.tex', '.latex']:
            return DocumentReader._read_text_file(file_path)
        elif suffix == '.docx':
            return DocumentReader._read_docx_file(file_path)
        elif suffix == '.pdf':
            return DocumentReader._read_pdf_file(file_path)
        else:
            # Try to read as text file anyway
            try:
                return DocumentReader._read_text_file(file_path)
            except Exception:
                raise ValueError(
                    f"Unsupported file format: {suffix}. "
                    f"Supported formats: .txt, .md, .docx, .pdf"
                )

    @staticmethod
    def _read_text_file(file_path: Path) -> str:
        """Read a plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def _read_docx_file(file_path: Path) -> str:
        """Read a Word document (.docx) file, including footnotes and endnotes."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required to read .docx files. "
                "Install it with: pip install python-docx"
            )

        doc = Document(file_path)

        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        # Extract footnotes and endnotes
        # Word stores these in separate XML parts that need special handling
        try:
            footnotes_text = DocumentReader._extract_docx_notes(file_path, 'footnotes')
            if footnotes_text:
                paragraphs.append("\nFootnotes:\n" + footnotes_text)
        except Exception:
            # If footnote extraction fails, continue without them
            pass

        try:
            endnotes_text = DocumentReader._extract_docx_notes(file_path, 'endnotes')
            if endnotes_text:
                paragraphs.append("\nEndnotes:\n" + endnotes_text)
        except Exception:
            # If endnote extraction fails, continue without them
            pass

        return '\n'.join(paragraphs)

    @staticmethod
    def _extract_docx_notes(file_path: Path, note_type: str) -> str:
        """
        Extract footnotes or endnotes from a .docx file.

        Args:
            file_path: Path to the .docx file
            note_type: Either 'footnotes' or 'endnotes'

        Returns:
            Extracted notes text
        """
        import zipfile
        from xml.etree import ElementTree as ET

        # .docx files are actually zip archives
        notes_text = []

        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Footnotes are stored in word/footnotes.xml
                # Endnotes are stored in word/endnotes.xml
                note_file = f'word/{note_type}.xml'

                if note_file not in docx_zip.namelist():
                    return ""

                # Read and parse the notes XML
                notes_xml = docx_zip.read(note_file)
                root = ET.fromstring(notes_xml)

                # Define the Word namespace
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }

                # Find all footnote/endnote elements
                for note in root.findall('.//w:footnote' if note_type == 'footnotes' else './/w:endnote', namespaces):
                    # Get the note ID
                    note_id = note.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

                    # Skip special notes (separator and continuation separator)
                    note_type_attr = note.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if note_type_attr in ['separator', 'continuationSeparator']:
                        continue

                    # Extract all text from this note
                    texts = []
                    for text_elem in note.findall('.//w:t', namespaces):
                        if text_elem.text:
                            texts.append(text_elem.text)

                    if texts:
                        note_text = ''.join(texts)
                        # Format as "note_id note_text"
                        notes_text.append(f"{note_id} {note_text}")

        except Exception:
            # If anything goes wrong, return empty string
            return ""

        return '\n'.join(notes_text)

    @staticmethod
    def _read_pdf_file(file_path: Path) -> str:
        """Read a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError(
                    "pypdf is required to read PDF files. "
                    "Install it with: pip install pypdf"
                )

        # Read PDF file
        reader = PdfReader(file_path)

        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return '\n'.join(text_parts)
